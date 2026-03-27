"""
Genereert Azure ML Stappenplan - Rister Pipeline v4.docx
Versie 4: cat_codes.json correctie + volledig C# API hoofdstuk
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

OUTPUT = "Azure ML Stappenplan - Rister Pipeline.docx"

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return h

def add_step(doc, nummer, titel, inhoud_func):
    add_heading(doc, f"Stap {nummer}: {titel}", level=2)
    inhoud_func(doc)
    doc.add_paragraph()

def code_block(doc, code_text):
    p = doc.add_paragraph()
    p.style = doc.styles["No Spacing"]
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F4FF")
    p._p.get_or_add_pPr().append(shading)
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(1 + level * 0.5)
    p.add_run(text)
    return p

def bold_intro(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p

def normal(doc, text):
    return doc.add_paragraph(text)


# ── Document aanmaken ────────────────────────────────────────────────────────
doc = Document()

# Paginamarges
section = doc.sections[0]
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Titelpagina
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Azure ML Stappenplan\nRister Pipeline")
tr.font.size = Pt(26)
tr.font.bold = True
tr.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Versie 4.0  |  " + datetime.date.today().strftime("%d %B %Y"))
sr.font.size = Pt(12)
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
desc = doc.add_paragraph()
desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
dr = desc.add_run(
    "Stap-voor-stap handleiding voor het inrichten van Azure ML,\n"
    "het draaien van de Rister ML-pipeline en het koppelen van de C# API."
)
dr.font.size = Pt(11)

doc.add_page_break()

# ── Inleiding ────────────────────────────────────────────────────────────────
add_heading(doc, "Overzicht", level=1)
normal(doc,
    "Dit document beschrijft alle stappen om de Rister ML-pipeline operationeel te maken in Azure. "
    "De pipeline traint een LightGBM regressor (tijdvoorspelling) en ranker (medewerkergeschiktheid), "
    "exporteert de modellen als ONNX, logt alles via MLflow en registreert de modellen in de "
    "Azure ML model registry. De C# API downloadt automatisch de nieuwste modellen bij opstart."
)

doc.add_paragraph()
bold_intro(doc, "Architectuuroverzicht:")
bullet(doc, "WerkExpert (lokaal, 26 Firebird DBs) → Parquet → Azure Blob Storage (Cool tier)")
bullet(doc, "Rister (PostgreSQL, Azure) → Azure ML Compute")
bullet(doc, "Azure ML Pipeline: rister-stap → merge-stap → training-stap")
bullet(doc, "MLflow tracking + model registry ingebouwd in Azure ML")
bullet(doc, "C# API: download bij opstart via MLflow REST API + DefaultAzureCredential")
bullet(doc, "Kwartaalschema: 1 jan / 1 apr / 1 jul / 1 okt om 02:00")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DEEL 1 – AZURE INRICHTEN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Deel 1 – Azure omgeving inrichten", level=1)

# Stap 1
def stap1(doc):
    normal(doc, "Maak een Azure ML workspace aan in de Azure Portal.")
    bullet(doc, "Resource group: rister-ml")
    bullet(doc, "Workspace naam: rister-aml")
    bullet(doc, "Regio: West Europe")
    bullet(doc, "Storage account, Key Vault en Container Registry worden automatisch aangemaakt.")
    doc.add_paragraph()
    normal(doc, "Na aanmaken noteer je de volgende waarden (nodig in stap 4):")
    code_block(doc,
        "Subscription ID : xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx\n"
        "Resource Group  : rister-ml\n"
        "Workspace naam  : rister-aml"
    )
add_step(doc, 1, "Azure ML Workspace aanmaken", stap1)

# Stap 2
def stap2(doc):
    normal(doc, "Maak een compute cluster aan voor de pipeline-stappen.")
    bullet(doc, "Naam: rister-cluster")
    bullet(doc, "VM-grootte: Standard_DS2_v2 (2 vCPU, 7 GB RAM)")
    bullet(doc, "Prioriteit: LowPriority (spot) — ~80% goedkoper")
    bullet(doc, "Minimum nodes: 0 (schaalt terug naar 0 bij inactiviteit)")
    bullet(doc, "Maximum nodes: 2")
    bullet(doc, "Idle scale-down: 60 seconden")
    doc.add_paragraph()
    normal(doc, "Via Azure CLI:")
    code_block(doc,
        "az ml compute create \\\n"
        "  --name rister-cluster \\\n"
        "  --type AmlCompute \\\n"
        "  --size Standard_DS2_v2 \\\n"
        "  --tier LowPriority \\\n"
        "  --min-instances 0 \\\n"
        "  --max-instances 2 \\\n"
        "  --idle-time-before-scale-down 60 \\\n"
        "  --workspace-name rister-aml \\\n"
        "  --resource-group rister-ml"
    )
    doc.add_paragraph()
    normal(doc, "Of via src/aml/deploy.py (doet dit automatisch bij eerste uitvoering).")
add_step(doc, 2, "Compute cluster aanmaken", stap2)

# Stap 3
def stap3(doc):
    normal(doc, "Maak een Azure Blob Storage container aan voor het WerkExpert Parquet-bestand.")
    bullet(doc, "Storage account: gebruik het account dat met de workspace is aangemaakt, of maak een nieuw account aan")
    bullet(doc, "Container naam: werkexpert-data")
    bullet(doc, "Tier: Cool (goedkoper voor zelden-gelezen bestanden)")
    doc.add_paragraph()
    normal(doc, "Noteer de connection string (Settings → Access keys → Connection string).")
    normal(doc, "Stel in als omgevingsvariabele:")
    code_block(doc, "AZURE_STORAGE_CONNECTION_STRING=DefaultEndpointsProtocol=https;AccountName=...;...")
    doc.add_paragraph()
    normal(doc,
        "werkexpert.py uploadt automatisch het gegenereerde Parquet-bestand naar deze container. "
        "Daarna leest de Azure ML pipeline het rechtstreeks via de blob URI."
    )
add_step(doc, 3, "Blob Storage inrichten voor WerkExpert", stap3)

# Stap 4
def stap4(doc):
    normal(doc, "Installeer de benodigde Python-pakketten en stel omgevingsvariabelen in.")
    code_block(doc,
        "pip install azure-ai-ml azure-identity mlflow\n"
        "pip install -r requirements.txt"
    )
    doc.add_paragraph()
    normal(doc, "Maak een .env bestand op basis van .env.example:")
    code_block(doc,
        "AML_SUBSCRIPTION_ID=xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx\n"
        "AML_RESOURCE_GROUP=rister-ml\n"
        "AML_WORKSPACE=rister-aml\n"
        "AZURE_STORAGE_CONNECTION_STRING=...\n"
        "MLFLOW_TRACKING_URI=azureml://westeurope.api.azureml.ms/mlflow/v1.0/subscriptions/..."
    )
    doc.add_paragraph()
    normal(doc, "Login met Azure CLI (voor lokaal gebruik):")
    code_block(doc, "az login")
add_step(doc, 4, "Lokale omgeving configureren", stap4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DEEL 2 – PIPELINE DRAAIEN
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Deel 2 – Pipeline draaien", level=1)

# Stap 5
def stap5(doc):
    normal(doc,
        "WerkExpert draait altijd lokaal omdat de 26 Firebird-databases alleen lokaal "
        "bereikbaar zijn (Docker container op Mac/Windows). Het resultaat wordt als "
        "Parquet-bestand geüpload naar Azure Blob Storage."
    )
    code_block(doc, "python src/dataprep/werkexpert.py")
    doc.add_paragraph()
    normal(doc, "Het script:")
    bullet(doc, "Verbindt met 26 Firebird DBs via FDB")
    bullet(doc, "Voert feature engineering uit (TF-IDF + SVD voor tekstkolommen, geodata)")
    bullet(doc, "Slaat op als Parquet (Snappy-compressie, ~80% kleiner dan CSV)")
    bullet(doc, "Uploadt naar Azure Blob Storage container werkexpert-data (Cool tier)")
    doc.add_paragraph()
    normal(doc, "Frequentie: handmatig of via lokale Prefect-flow (kwartaallijks aanbevolen).")
add_step(doc, 5, "WerkExpert lokaal draaien en uploaden", stap5)

# Stap 6
def stap6(doc):
    normal(doc, "Submit de Azure ML pipeline. Dit start drie stappen op het compute cluster:")
    bullet(doc, "Stap 1 – rister: haalt data op uit PostgreSQL (Azure), slaat op als CSV")
    bullet(doc, "Stap 2 – merge: combineert rister.csv + werkexpert.parquet uit Blob Storage")
    bullet(doc, "Stap 3 – training: traint LightGBM regressor + ranker, exporteert ONNX, logt naar MLflow")
    doc.add_paragraph()
    code_block(doc, "python src/aml/pipeline.py")
    doc.add_paragraph()
    normal(doc, "Voortgang volgen in Azure ML Studio → Jobs.")
add_step(doc, 6, "Azure ML pipeline starten", stap6)

# Stap 7
def stap7(doc):
    normal(doc,
        "Het kwartaalschema wordt eenmalig ingesteld. Daarna start de pipeline automatisch "
        "op 1 januari, 1 april, 1 juli en 1 oktober om 02:00 UTC."
    )
    code_block(doc, "python src/aml/deploy.py")
    doc.add_paragraph()
    normal(doc, "Wat deploy.py doet:")
    bullet(doc, "Maakt het compute cluster aan als het nog niet bestaat")
    bullet(doc, "Registreert de pipeline als een herbruikbaar component")
    bullet(doc, "Maakt een kwartaalschedule aan: cron \"0 2 1 1,4,7,10 *\"")
    doc.add_paragraph()
    bold_intro(doc, "Schedule aanpassen in Azure ML Studio:")
    normal(doc,
        "Ga naar Azure ML Studio → Jobs → Schedules. Je kunt de CRON-expressie "
        "direct aanpassen in de Studio. Let op: de wijziging in de Studio wordt "
        "niet automatisch teruggeschreven naar deploy.py. Wil je dit in code bijhouden, "
        "pas dan ook de CRON_EXPRESSION variabele in deploy.py aan."
    )
add_step(doc, 7, "Kwartaalschedule instellen", stap7)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DEEL 3 – MODEL REGISTRY
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Deel 3 – Model registry", level=1)

# Stap 8
def stap8(doc):
    normal(doc,
        "Na elke succesvolle training vergelijkt het trainingsscript de nieuwe modellen met de "
        "huidige productieversie. Als de nieuwe modellen beter presteren, worden ze automatisch "
        "geregistreerd in de Azure ML model registry."
    )
    doc.add_paragraph()
    normal(doc, "Geregistreerde modelnamen:")
    bullet(doc, "rister-lgbm-regressor  (tijdvoorspelling)")
    bullet(doc, "rister-lgbm-ranker     (medewerkergeschiktheid)")
    doc.add_paragraph()
    normal(doc, "Controleer de registry via Azure ML Studio → Models, of via CLI:")
    code_block(doc,
        "az ml model list --workspace-name rister-aml --resource-group rister-ml\n"
        "az ml model show --name rister-lgbm-regressor --workspace-name rister-aml --resource-group rister-ml"
    )
    doc.add_paragraph()
    normal(doc, "Elk geregistreerd model bevat de volgende artifacts:")
    bullet(doc, "regressor.onnx  /  ranker.onnx  — ONNX-modelbestanden voor inferentie")
    bullet(doc, "models/cat_codes.json  — categorische label-encodings (Python-formaat)")
    bullet(doc, "models/scaler.json  — mean/scale voor numerieke feature-normalisatie")
    bullet(doc, "models/metadata.json  — trainingsmetadata en feature-namen")
add_step(doc, 8, "Modellen bekijken in de registry", stap8)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DEEL 4 – C# API KOPPELING
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Deel 4 – C# API koppelen aan Azure ML", level=1)

normal(doc,
    "De C# API downloadt automatisch de nieuwste modelbestanden uit de Azure ML model registry "
    "bij elke opstart. Dit verloopt volledig via C#-code zonder externe scripts. "
    "De authenticatie gebruikt DefaultAzureCredential: lokaal via az login, "
    "in productie via Managed Identity."
)
doc.add_paragraph()

# Stap 9
def stap9(doc):
    normal(doc, "Voeg de volgende sectie toe aan appsettings.json (of stel in als omgevingsvariabelen):")
    code_block(doc,
        "\"ML\": {\n"
        "  \"RegressorPath\":  \"ml/regressor.onnx\",\n"
        "  \"RankerPath\":     \"ml/ranker.onnx\",\n"
        "  \"EncodersPath\":   \"ml/cat_codes.json\",\n"
        "  \"ScalerPath\":     \"ml/scaler.json\",\n"
        "  \"MlflowUri\":      \"https://westeurope.api.azureml.ms/mlflow/v1.0/subscriptions/\n"
        "                      {SUBSCRIPTION_ID}/resourceGroups/{RG}/providers/\n"
        "                      Microsoft.MachineLearningServices/workspaces/{WS}\",\n"
        "  \"RegressorName\":  \"rister-lgbm-regressor\",\n"
        "  \"RankerName\":     \"rister-lgbm-ranker\"\n"
        "}"
    )
    doc.add_paragraph()
    normal(doc,
        "MlflowUri leeg laten om Azure ML-download over te slaan en de lokale bestanden in ml/ te gebruiken. "
        "Dit is handig voor lokale ontwikkeling."
    )
    doc.add_paragraph()
    bold_intro(doc, "Bestandsnamen in ml/:")
    bullet(doc, "ml/regressor.onnx  — tijdvoorspellingsmodel")
    bullet(doc, "ml/ranker.onnx     — medewerkergeschiktheidsmodel")
    bullet(doc, "ml/cat_codes.json  — categorische encodings (formaat: {\"kolom\": {\"waarde\": int_code}})")
    bullet(doc, "ml/scaler.json     — numerieke normalisatie (mean + scale arrays)")
add_step(doc, 9, "appsettings.json configureren", stap9)

# Stap 10
def stap10(doc):
    normal(doc,
        "De API gebruikt drie ML-klassen in Rister.Manure.Infrastructure/ML/. "
        "Hieronder staat de werking per klasse."
    )
    doc.add_paragraph()

    bold_intro(doc, "MlOptions.cs — configuratieklasse")
    normal(doc, "Leest de ML-sectie uit appsettings.json via IOptions<MlOptions>.")
    code_block(doc,
        "public sealed class MlOptions\n"
        "{\n"
        "    public const string Section = \"ML\";\n"
        "    public string RegressorPath  { get; init; } = \"ml/regressor.onnx\";\n"
        "    public string RankerPath     { get; init; } = \"ml/ranker.onnx\";\n"
        "    public string EncodersPath   { get; init; } = \"ml/cat_codes.json\";\n"
        "    public string ScalerPath     { get; init; } = \"ml/scaler.json\";\n"
        "    public string MlflowUri      { get; init; } = string.Empty;\n"
        "    public string RegressorName  { get; init; } = \"rister-lgbm-regressor\";\n"
        "    public string RankerName     { get; init; } = \"rister-lgbm-ranker\";\n"
        "}"
    )
    doc.add_paragraph()

    bold_intro(doc, "AzureMLModelDownloader.cs — download bij opstart")
    normal(doc,
        "Downloadt de nieuwste modellen vóór het laden van de ONNX-sessies. "
        "Bij een fout wordt teruggevallen op de bestaande lokale bestanden."
    )
    bullet(doc, "Authenticatie: DefaultAzureCredential (az login lokaal, Managed Identity in Azure)")
    bullet(doc, "Haalt run_id op via: POST .../registered-models/get-latest-versions")
    bullet(doc, "Downloadt artifacts via: GET .../artifacts/get?run_id={id}&path={path}")
    bullet(doc, "Downloads: regressor.onnx, ranker.onnx, models/cat_codes.json")
    doc.add_paragraph()

    bold_intro(doc, "FeatureBuilder.cs — feature-engineering in C#")
    normal(doc,
        "Zet een PredictionInput om naar ModelFeatures (float[36]) voor ONNX-inferentie. "
        "De volgorde is exact gelijk aan de Python-training: "
        "categorisch (index 0–10) + numeriek genormaliseerd (index 11–35)."
    )
    normal(doc, "cat_codes.json formaat (Python-formaat, direct leesbaar door C#):")
    code_block(doc,
        "// cat_codes.json\n"
        "{\n"
        "  \"URENVERANTW_MEDID\": { \"EMP001\": 0, \"EMP002\": 1, ... },\n"
        "  \"BEWERKING_ID\":      { \"BW100\":  0, \"BW101\":  1, ... },\n"
        "  ...\n"
        "}"
    )
    normal(doc, "C# leest dit direct als Dictionary<string, Dictionary<string, int>>:")
    code_block(doc,
        "this._encoders = JsonSerializer.Deserialize<\n"
        "    Dictionary<string, Dictionary<string, int>>>(\n"
        "    File.ReadAllText(encodersPath))!;"
    )
    doc.add_paragraph()

    bold_intro(doc, "DependencyInjection.cs — volgorde van registratie")
    normal(doc, "De download-volgorde is bewust synchroon om een race-condition te voorkomen:")
    code_block(doc,
        "// 1. ML-configuratie laden\n"
        "builder.Services.Configure<MlOptions>(builder.Configuration.GetSection(\"ML\"));\n\n"
        "// 2. Modellen downloaden VÓÓR het laden van de ONNX-sessies\n"
        "DownloadModelsIfConfigured(builder);\n\n"
        "// 3. ONNX singletons laden vanuit (bijgewerkte) lokale bestanden\n"
        "builder.Services.AddSingleton<IFeatureBuilder>(\n"
        "    _ => new FeatureBuilder(ml.EncodersPath, ml.ScalerPath));\n"
        "builder.Services.AddSingleton<IOnnxRunner>(\n"
        "    _ => new OnnxRunner(ml.RegressorPath, ml.RankerPath));"
    )
add_step(doc, 10, "C# API ML-klassen", stap10)

# Stap 11
def stap11(doc):
    normal(doc, "Voor Azure Container Apps / App Service: zet Managed Identity aan.")
    bullet(doc, "Ga naar de Container App → Identity → System assigned → zet op On")
    bullet(doc, "Kopieer de Object (principal) ID")
    doc.add_paragraph()
    normal(doc, "Ken de rol toe zodat de API de model registry kan lezen:")
    code_block(doc,
        "az role assignment create \\\n"
        "  --assignee <PRINCIPAL_ID> \\\n"
        "  --role \"AzureML Data Scientist\" \\\n"
        "  --scope /subscriptions/<SUB>/resourceGroups/rister-ml/providers/\\\n"
        "          Microsoft.MachineLearningServices/workspaces/rister-aml"
    )
    doc.add_paragraph()
    normal(doc,
        "Lokaal werkt DefaultAzureCredential automatisch via az login. "
        "Geen extra configuratie nodig."
    )
    doc.add_paragraph()
    bold_intro(doc, "NuGet-pakket vereist:")
    code_block(doc, "<PackageReference Include=\"Azure.Identity\" Version=\"1.13.2\" />")
add_step(doc, 11, "Managed Identity instellen voor de API", stap11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# DEEL 5 – KOSTEN & BEHEER
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Deel 5 – Kostenoptimalisatie en beheer", level=1)

# Stap 12
def stap12(doc):
    normal(doc, "De volgende maatregelen zijn geïmplementeerd om Azure-kosten te minimaliseren:")
    doc.add_paragraph()

    bold_intro(doc, "Compute:")
    bullet(doc, "LowPriority (spot) VM's: ~80% goedkoper dan dedicated")
    bullet(doc, "Min nodes = 0: cluster schaalt volledig af bij inactiviteit")
    bullet(doc, "Idle scale-down na 60 seconden (standaard is 120s)")
    bullet(doc, "VM-grootte Standard_DS2_v2: voldoende voor LightGBM, geen GPU nodig")
    doc.add_paragraph()

    bold_intro(doc, "Storage:")
    bullet(doc, "WerkExpert Parquet (Snappy): ~80% kleiner dan CSV")
    bullet(doc, "Azure Blob Storage Cool tier: ~50% goedkoper dan Hot voor zelden-gelezen data")
    bullet(doc, "ONNX-modellen zijn klein (<10 MB), geen speciale opslag nodig")
    doc.add_paragraph()

    bold_intro(doc, "Uitvoerfrequentie:")
    bullet(doc, "Kwartaalschedule (4x per jaar) i.p.v. wekelijks: ~92% minder compute-uren")
    bullet(doc, "WerkExpert-data verandert nauwelijks; Rister-data groeit automatisch")
    doc.add_paragraph()

    bold_intro(doc, "Geschatte maandelijkse kosten (kwartaallijks draaien):")
    bullet(doc, "Compute (LowPriority, ~2u per kwartaal): €0,05–0,10 per run")
    bullet(doc, "Blob Storage (Cool tier, <1 GB): €0,01 per maand")
    bullet(doc, "Azure ML workspace (gratis tier beschikbaar): €0")
    bullet(doc, "Totaal: <€1 per maand")
add_step(doc, 12, "Kostenoptimalisatie", stap12)

# Stap 13
def stap13(doc):
    normal(doc, "Controleer de pipeline-runs en modelkwaliteit via:")
    doc.add_paragraph()

    bold_intro(doc, "Azure ML Studio (ml.azure.com):")
    bullet(doc, "Jobs → bekijk alle pipeline-runs en logs")
    bullet(doc, "Models → bekijk geregistreerde versies van regressor en ranker")
    bullet(doc, "Experiments → bekijk MLflow metrics per run (MAE, NDCG, etc.)")
    bullet(doc, "Schedules → pas de kwartaalplanning aan")
    doc.add_paragraph()

    bold_intro(doc, "Lokale Prefect monitoring (optioneel):")
    code_block(doc,
        "prefect server start         # start lokale Prefect UI op http://localhost:4200\n"
        "python src/flows/pipeline_flow.py   # handmatig starten"
    )
    doc.add_paragraph()

    bold_intro(doc, "API startup log (bij elke herstart):")
    normal(doc,
        "De C# API logt bij opstart of de download geslaagd is. Controleer de applicatielogs "
        "op berichten van AzureMLModelDownloader:"
    )
    code_block(doc,
        "[INF] Azure ML: ophalen laatste modelversies...\n"
        "[INF] Model 'rister-lgbm-regressor' versie 3 gevonden (run: abc123)\n"
        "[INF] Gedownload: regressor.onnx → ml/regressor.onnx (2048 KB)\n"
        "[INF] Gedownload: ranker.onnx → ml/ranker.onnx (1536 KB)\n"
        "[INF] Gedownload: models/cat_codes.json → ml/cat_codes.json (12 KB)\n"
        "[INF] Azure ML: modelbestanden bijgewerkt."
    )
    normal(doc, "Bij een mislukte download:")
    code_block(doc,
        "[WRN] Azure ML download mislukt — bestaande lokale bestanden worden gebruikt."
    )
add_step(doc, 13, "Monitoring en beheer", stap13)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# BIJLAGE – BESTANDSSTRUCTUUR
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Bijlage – Bestandsstructuur", level=1)

normal(doc, "Python pipeline:")
code_block(doc,
    "src/\n"
    "├── config.py                    # Paden, DB-verbindingen, MLflow URI\n"
    "├── dataprep/\n"
    "│   ├── rister.py                # PostgreSQL → data/processed/rister.csv\n"
    "│   ├── werkexpert.py            # 26 Firebird DBs → werkexpert.parquet + Blob upload\n"
    "│   └── merge.py                 # rister.csv + werkexpert.parquet → gecombineerd.csv\n"
    "├── train/\n"
    "│   └── train_lightgbm.py        # LightGBM → ONNX + MLflow + model registry\n"
    "├── aml/\n"
    "│   ├── pipeline.py              # Azure ML pipeline definitie + submit\n"
    "│   └── deploy.py                # Compute + kwartaalschedule aanmaken\n"
    "├── flows/\n"
    "│   ├── pipeline_flow.py         # Prefect flow (lokale orchestratie)\n"
    "│   └── deploy.py                # Prefect deployment\n"
    "└── pipeline.py                  # Lokale pipeline runner (alle stappen)"
)
doc.add_paragraph()

normal(doc, "C# API (Rister.Manure.Infrastructure/ML/):")
code_block(doc,
    "ML/\n"
    "├── MlOptions.cs                 # Configuratieklasse (appsettings.json → ML-sectie)\n"
    "├── AzureMLModelDownloader.cs    # Download bij opstart via MLflow REST API\n"
    "├── FeatureBuilder.cs            # PredictionInput → float[36] voor ONNX\n"
    "├── OnnxRunner.cs                # ONNX-inferentie (regressor + ranker)\n"
    "└── ScalerData.cs                # Deserializatie van scaler.json\n"
    "\n"
    "ml/  (runtime, niet in git)\n"
    "├── regressor.onnx\n"
    "├── ranker.onnx\n"
    "├── cat_codes.json               # {\"kolom\": {\"waarde\": int_code}}\n"
    "└── scaler.json                  # {\"mean\": [...], \"scale\": [...]}"
)
doc.add_paragraph()

normal(doc, "Model artifacts in Azure ML registry (per run):")
code_block(doc,
    "regressor.onnx\n"
    "ranker.onnx\n"
    "models/\n"
    "├── cat_codes.json\n"
    "├── scaler.json\n"
    "└── metadata.json"
)

# ══════════════════════════════════════════════════════════════════════════════
# VERSIEGESCHIEDENIS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
add_heading(doc, "Versiegeschiedenis", level=1)

versies = [
    ("4.0", "2026-03-27", "cat_codes.json fix (encoders.json was onjuist), C# API volledig gedocumenteerd"),
    ("3.0", "2026-03",    "C# API koppeling aan Azure ML model registry toegevoegd (puur C#, geen scripts)"),
    ("2.0", "2026-03",    "WerkExpert Parquet + Blob Storage, kostenoptimalisatie, Firebird-splitsing"),
    ("1.0", "2026-02",    "Initiële versie: Azure ML pipeline, kwartaalschedule, Prefect"),
]
for versie, datum, omschrijving in versies:
    p = doc.add_paragraph()
    p.add_run(f"v{versie}  ({datum})  ").bold = True
    p.add_run(omschrijving)

# ── Opslaan ──────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"Document opgeslagen als: {OUTPUT}")
